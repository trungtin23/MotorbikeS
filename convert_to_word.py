from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import datetime

def add_heading_with_formatting(doc, text, level):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Set font
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        if level == 0:
            run.font.size = Pt(18)
        elif level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    
    return heading

def add_paragraph_with_formatting(doc, text, bold=False, italic=False):
    """Add a paragraph with formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def add_bullet_point(doc, text):
    """Add a bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def create_report():
    # Create document
    doc = Document()
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title page
    title = doc.add_heading('BÁO CÁO DỰ ÁN CỬA HÀNG XE MÁY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(20)
        run.bold = True
    
    subtitle = doc.add_heading('(MOTORBIKE STORE E-COMMERCE PLATFORM)', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.italic = True
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Current date
    current_date = datetime.datetime.now().strftime("%d/%m/%Y")
    date_para = doc.add_paragraph(f"Ngày báo cáo: {current_date}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    # Page break
    doc.add_page_break()
    
    # I. TỔNG QUAN DỰ ÁN
    add_heading_with_formatting(doc, 'I. TỔNG QUAN DỰ ÁN', 1)
    
    add_heading_with_formatting(doc, '1.1 Giới thiệu', 2)
    add_paragraph_with_formatting(doc, 'Dự án "Cửa hàng xe máy" là một hệ thống thương mại điện tử hoàn chỉnh được phát triển để quản lý và bán hàng trực tuyến các sản phẩm xe máy. Hệ thống được xây dựng với kiến trúc Full-Stack, sử dụng các công nghệ hiện đại để đảm bảo hiệu suất cao và trải nghiệm người dùng tốt nhất.')
    
    add_heading_with_formatting(doc, '1.2 Mục tiêu dự án', 2)
    add_bullet_point(doc, 'Xây dựng hệ thống bán hàng trực tuyến chuyên về xe máy')
    add_bullet_point(doc, 'Quản lý danh mục sản phẩm, thương hiệu và dòng xe')
    add_bullet_point(doc, 'Hỗ trợ đặt hàng, thanh toán trực tuyến')
    add_bullet_point(doc, 'Quản lý tài khoản người dùng và nhân viên')
    add_bullet_point(doc, 'Hệ thống đánh giá và bình luận sản phẩm')
    add_bullet_point(doc, 'Tích hợp thanh toán VNPay')
    add_bullet_point(doc, 'Hệ thống xác thực email và bảo mật')
    
    # II. KIẾN TRÚC VÀ CÔNG NGHỆ SỬ DỤNG
    add_heading_with_formatting(doc, 'II. KIẾN TRÚC VÀ CÔNG NGHỆ SỬ DỤNG', 1)
    
    add_heading_with_formatting(doc, '2.1 Kiến trúc tổng thể', 2)
    add_paragraph_with_formatting(doc, 'Dự án được xây dựng theo mô hình 3-tier architecture:')
    add_bullet_point(doc, 'Presentation Layer: Frontend React.js')
    add_bullet_point(doc, 'Business Logic Layer: Backend Spring Boot')
    add_bullet_point(doc, 'Data Access Layer: MySQL Database')
    
    add_heading_with_formatting(doc, '2.2 Công nghệ Backend', 2)
    add_heading_with_formatting(doc, '2.2.1 Framework và Dependencies chính', 3)
    add_bullet_point(doc, 'Spring Boot 3.4.4: Framework chính cho phát triển backend')
    add_bullet_point(doc, 'Java 17: Ngôn ngữ lập trình chính')
    add_bullet_point(doc, 'Spring Data JPA: Quản lý dữ liệu và ORM')
    add_bullet_point(doc, 'Spring Security: Bảo mật và xác thực')
    add_bullet_point(doc, 'Spring Boot Web: Phát triển REST API')
    add_bullet_point(doc, 'MySQL Connector: Kết nối cơ sở dữ liệu')
    
    add_heading_with_formatting(doc, '2.2.2 Libraries và Tools', 3)
    add_bullet_point(doc, 'JWT (jsonwebtoken 0.12.6): Xác thực và phân quyền')
    add_bullet_point(doc, 'ModelMapper 3.0.0: Mapping giữa entities và DTOs')
    add_bullet_point(doc, 'Lombok: Giảm boilerplate code')
    add_bullet_point(doc, 'Spring Boot Mail: Gửi email xác thực')
    add_bullet_point(doc, 'Spring Boot DevTools: Hỗ trợ phát triển')
    
    add_heading_with_formatting(doc, '2.2.3 Database', 3)
    add_bullet_point(doc, 'MySQL: Hệ quản trị cơ sở dữ liệu chính')
    add_bullet_point(doc, 'Hibernate: ORM framework')
    add_bullet_point(doc, 'Connection Pool: Quản lý kết nối database')
    
    add_heading_with_formatting(doc, '2.3 Công nghệ Frontend', 2)
    add_heading_with_formatting(doc, '2.3.1 Framework và Libraries', 3)
    add_bullet_point(doc, 'React 19.1.0: Framework JavaScript chính')
    add_bullet_point(doc, 'Vite 6.2.0: Build tool và development server')
    add_bullet_point(doc, 'React Router DOM 7.6.0: Routing navigation')
    add_bullet_point(doc, 'Axios 1.9.0: HTTP client cho API calls')
    
    add_heading_with_formatting(doc, '2.3.2 UI/UX Libraries', 3)
    add_bullet_point(doc, 'Material-UI (MUI) 7.1.0: Component library')
    add_bullet_point(doc, 'TailwindCSS 4.1.6: Utility-first CSS framework')
    add_bullet_point(doc, 'Lucide React 0.509.0: Icon library')
    add_bullet_point(doc, 'Emotion React/Styled: CSS-in-JS styling')
    
    # III. CẤU TRÚC DỮ LIỆU
    add_heading_with_formatting(doc, 'III. CẤU TRÚC DỮ LIỆU', 1)
    
    add_heading_with_formatting(doc, '3.1 Database Schema', 2)
    add_paragraph_with_formatting(doc, 'Hệ thống sử dụng MySQL với tên database: motobike_web')
    
    add_heading_with_formatting(doc, '3.2 Các Entity chính', 2)
    
    add_heading_with_formatting(doc, '3.2.1 Account (Tài khoản)', 3)
    account_fields = [
        'id (Integer): Primary key',
        'username (String): Tên đăng nhập',
        'password (String): Mật khẩu đã mã hóa',
        'email (String): Email người dùng',
        'phone (String): Số điện thoại',
        'name (String): Họ và tên',
        'role (String): Vai trò (1=User, 2=Admin)',
        'status (String): Trạng thái (ACTIVE/PENDING)',
        'securityCode (String): Mã xác thực',
        'address (String): Địa chỉ',
        'dob (Instant): Ngày sinh',
        'avatar (String): Ảnh đại diện',
        'created (Instant): Ngày tạo'
    ]
    for field in account_fields:
        add_bullet_point(doc, field)
    
    add_heading_with_formatting(doc, '3.2.2 Product (Sản phẩm)', 3)
    product_fields = [
        'id (Integer): Primary key',
        'name (String): Tên sản phẩm',
        'description (String): Mô tả',
        'price (Double): Giá',
        'avatar (String): Ảnh đại diện',
        'weight (String): Khối lượng',
        'size (String): Kích thước',
        'petrolCapacity (String): Dung tích bình xăng',
        'saddleHeight (String): Chiều cao yên',
        'wheelSize (String): Kích thước bánh xe',
        'beforeFork/afterFork (String): Hệ thống treo',
        'maxiumCapacity (String): Công suất tối đa',
        'oilCapacity (String): Dung tích nhớt',
        'fuelConsumption (String): Mức tiêu thụ nhiên liệu',
        'cylinderCapacity (String): Dung tích xi-lanh',
        'maxiumMoment (String): Mô-men xoắn tối đa',
        'compressionRatio (String): Tỷ số nén',
        'engieType (String): Loại động cơ',
        'brandId (Integer): Thương hiệu',
        'motolineId (Integer): Dòng xe'
    ]
    for field in product_fields:
        add_bullet_point(doc, field)
    
    # IV. CHỨC NĂNG CHÍNH
    add_heading_with_formatting(doc, 'IV. CHỨC NĂNG CHÍNH', 1)
    
    add_heading_with_formatting(doc, '4.1 Hệ thống Authentication & Authorization', 2)
    
    add_heading_with_formatting(doc, '4.1.1 Đăng ký tài khoản', 3)
    add_bullet_point(doc, 'Validation dữ liệu đầu vào')
    add_bullet_point(doc, 'Mã hóa mật khẩu với BCrypt')
    add_bullet_point(doc, 'Gửi email xác thực')
    add_bullet_point(doc, 'Trạng thái tài khoản PENDING cho đến khi xác thực')
    
    add_heading_with_formatting(doc, '4.1.2 Đăng nhập', 3)
    add_bullet_point(doc, 'Xác thực username/password')
    add_bullet_point(doc, 'Tạo JWT token với thông tin user và role')
    add_bullet_point(doc, 'Kiểm tra trạng thái tài khoản ACTIVE')
    
    add_heading_with_formatting(doc, '4.1.3 Quên mật khẩu', 3)
    add_bullet_point(doc, 'Gửi link reset password qua email')
    add_bullet_point(doc, 'Token có thời hạn cho bảo mật')
    add_bullet_point(doc, 'Đặt lại mật khẩu mới')
    
    add_heading_with_formatting(doc, '4.2 Quản lý sản phẩm', 2)
    add_bullet_point(doc, 'CRUD operations cho sản phẩm')
    add_bullet_point(doc, 'Quản lý thương hiệu và dòng xe')
    add_bullet_point(doc, 'Upload và quản lý hình ảnh')
    add_bullet_point(doc, 'Quản lý phiên bản và màu sắc sản phẩm')
    
    add_heading_with_formatting(doc, '4.3 Hệ thống đặt hàng', 2)
    add_bullet_point(doc, 'Thêm sản phẩm vào giỏ hàng')
    add_bullet_point(doc, 'Quản lý đơn hàng')
    add_bullet_point(doc, 'Tính toán tổng tiền')
    add_bullet_point(doc, 'Quản lý trạng thái đơn hàng')
    
    add_heading_with_formatting(doc, '4.4 Thanh toán', 2)
    add_bullet_point(doc, 'Tích hợp VNPay Gateway')
    add_bullet_point(doc, 'Hỗ trợ nhiều phương thức thanh toán')
    add_bullet_point(doc, 'Xử lý callback từ cổng thanh toán')
    
    # V. BẢO MẬT VÀ XÁC THỰC
    add_heading_with_formatting(doc, 'V. BẢO MẬT VÀ XÁC THỰC', 1)
    
    add_heading_with_formatting(doc, '5.1 Bảo mật Backend', 2)
    add_bullet_point(doc, 'Spring Security: Framework bảo mật chính')
    add_bullet_point(doc, 'JWT Authentication: Stateless authentication')
    add_bullet_point(doc, 'Password Encryption: BCrypt hashing')
    add_bullet_point(doc, 'CORS Configuration: Kiểm soát cross-origin requests')
    
    add_heading_with_formatting(doc, '5.2 Validation và Sanitization', 2)
    add_bullet_point(doc, 'Input validation cho all endpoints')
    add_bullet_point(doc, 'SQL Injection prevention với JPA')
    add_bullet_point(doc, 'XSS protection')
    
    # VI. TÍCH HỢP THIRD-PARTY
    add_heading_with_formatting(doc, 'VI. TÍCH HỢP THIRD-PARTY', 1)
    
    add_heading_with_formatting(doc, '6.1 VNPay Payment Gateway', 2)
    add_bullet_point(doc, 'Terminal ID: JRNO6CET')
    add_bullet_point(doc, 'Sandbox Environment: Testing payments')
    add_bullet_point(doc, 'Return URL Handling: Xử lý kết quả thanh toán')
    add_bullet_point(doc, 'Security: Hash validation')
    
    add_heading_with_formatting(doc, '6.2 Email Service', 2)
    add_bullet_point(doc, 'Gmail SMTP: smtp.gmail.com:587')
    add_bullet_point(doc, 'HTML Email Templates: Rich email content')
    add_bullet_point(doc, 'Verification Links: Secure token-based verification')
    
    # VII. CẤU HÌNH VÀ DEPLOYMENT
    add_heading_with_formatting(doc, 'VII. CẤU HÌNH VÀ DEPLOYMENT', 1)
    
    add_heading_with_formatting(doc, '7.1 Backend Configuration', 2)
    config_items = [
        'Server Port: 8080',
        'Database: MySQL (motobike_web)',
        'JPA: Hibernate with MySQL dialect',
        'Email: Gmail SMTP',
        'VNPay: Sandbox environment'
    ]
    for item in config_items:
        add_bullet_point(doc, item)
    
    add_heading_with_formatting(doc, '7.2 Frontend Configuration', 2)
    frontend_config = [
        'Development Server: Vite',
        'Port: 5173 (default)',
        'Build Tool: Vite',
        'Package Manager: npm'
    ]
    for item in frontend_config:
        add_bullet_point(doc, item)
    
    # VIII. KẾT LUẬN
    add_heading_with_formatting(doc, 'VIII. KẾT LUẬN', 1)
    
    add_paragraph_with_formatting(doc, 'Dự án "Cửa hàng xe máy" đã được phát triển thành công với kiến trúc hiện đại và các tính năng đầy đủ cho một hệ thống thương mại điện tử. Việc sử dụng Spring Boot cho backend và React cho frontend đảm bảo khả năng mở rộng và bảo trì tốt trong tương lai.')
    
    add_heading_with_formatting(doc, 'Ưu điểm của dự án:', 2)
    advantages = [
        'Kiến trúc rõ ràng: Separation of concerns tốt',
        'Bảo mật cao: JWT + Spring Security',
        'UX/UI hiện đại: Material-UI + TailwindCSS',
        'Tích hợp thanh toán: VNPay gateway',
        'Email verification: Bảo mật tài khoản',
        'Responsive design: Hỗ trợ mobile'
    ]
    for advantage in advantages:
        add_bullet_point(doc, advantage)
    
    add_heading_with_formatting(doc, 'Công nghệ sử dụng đáng chú ý:', 2)
    technologies = [
        'Java 17 + Spring Boot 3.4.4',
        'React 19.1.0 + Vite 6.2.0',
        'MySQL + JPA/Hibernate',
        'JWT Authentication',
        'Material-UI + TailwindCSS',
        'VNPay Integration'
    ]
    for tech in technologies:
        add_bullet_point(doc, tech)
    
    add_paragraph_with_formatting(doc, 'Dự án này thể hiện sự kết hợp tốt giữa các công nghệ backend và frontend hiện đại, tạo ra một sản phẩm hoàn chỉnh và có thể triển khai thực tế.')
    
    # Save document
    doc.save('BÁO_CÁO_DỰ_ÁN_CỬA_HÀNG_XE_MÁY.docx')
    print("Đã tạo file Word báo cáo thành công: BÁO_CÁO_DỰ_ÁN_CỬA_HÀNG_XE_MÁY.docx")

if __name__ == "__main__":
    create_report() 